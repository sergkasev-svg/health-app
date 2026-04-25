"""
Извлечение текста из загруженных документов (PDF, изображения) для анализа и отчёта.
OCR встроен в приложение (EasyOCR) — пользователю ничего дополнительно ставить не нужно.
При отсутствии EasyOCR используется Tesseract, если найден в системе.
"""
from __future__ import annotations

import logging
from pathlib import Path
from typing import Optional

logger = logging.getLogger(__name__)

MIN_PDF_TEXT_FOR_OCR = 200
MAX_EXTRACTED_TEXT_LENGTH = 100_000

# Кэш ридера EasyOCR (ленивая инициализация)
_easyocr_reader = None


def _normalize_text(text: Optional[str]) -> str:
    """Безопасная нормализация извлечённого текста."""
    if not text:
        return ""
    normalized = str(text).replace("\x00", " ")
    normalized = normalized.replace("\r\n", "\n").replace("\r", "\n")
    lines = [line.strip() for line in normalized.split("\n")]
    normalized = "\n".join(line for line in lines if line)
    normalized = normalized.strip()
    if len(normalized) > MAX_EXTRACTED_TEXT_LENGTH:
        normalized = normalized[:MAX_EXTRACTED_TEXT_LENGTH].strip()
    return normalized


def _get_easyocr_reader():
    """Ленивая инициализация EasyOCR (ru+en). При первом запуске скачивает модели (~100 МБ)."""
    global _easyocr_reader
    if _easyocr_reader is not None:
        return _easyocr_reader
    try:
        import easyocr

        logger.info("Loading EasyOCR (ru, en) for built-in OCR...")
        _easyocr_reader = easyocr.Reader(["ru", "en"], gpu=False, verbose=False)
        logger.info("EasyOCR ready.")
        return _easyocr_reader
    except ImportError:
        logger.warning("EasyOCR not installed. Run: pip install easyocr (recommended for scans)")
        return None
    except Exception as e:
        logger.warning("EasyOCR init failed: %s", e)
        return None


def _ocr_with_easyocr(image_or_path) -> str:
    """Распознавание текста через EasyOCR. Принимает Path, str или PIL Image."""
    reader = _get_easyocr_reader()
    if not reader:
        return ""
    try:
        import numpy as np

        if isinstance(image_or_path, (str, Path)):
            result = reader.readtext(str(image_or_path))
        else:
            img = image_or_path
            if hasattr(img, "mode") and img.mode != "RGB":
                img = img.convert("RGB")
            arr = np.array(img)
            result = reader.readtext(arr)

        lines = [item[1] for item in result if len(item) > 1 and str(item[1]).strip()]
        return _normalize_text("\n".join(lines))
    except Exception as e:
        logger.warning("EasyOCR readtext error: %s", e)
        return ""


def _tesseract_cmd_windows() -> None:
    """Подсказка пути к Tesseract на Windows, если не в PATH."""
    try:
        import pytesseract

        for path in (
            r"C:\Program Files\Tesseract-OCR\tesseract.exe",
            r"C:\Program Files (x86)\Tesseract-OCR\tesseract.exe",
        ):
            if Path(path).exists():
                pytesseract.pytesseract.tesseract_cmd = path
                return
    except Exception:
        return


def _ocr_with_tesseract(pil_img) -> str:
    """Резервный OCR через Tesseract (если установлен в системе)."""
    _tesseract_cmd_windows()
    try:
        import pytesseract
    except ImportError:
        return ""

    try:
        text = pytesseract.image_to_string(pil_img, lang="rus+eng")
        return _normalize_text(text)
    except Exception as e:
        error_text = str(e).lower()
        if "not found" in error_text or "not in your path" in error_text:
            logger.debug("Tesseract not available: %s", e)
        else:
            logger.warning("Tesseract OCR error: %s", e)
        return ""


def _run_ocr(pil_image_or_path) -> str:
    """Единая точка OCR: сначала EasyOCR, при неудаче — Tesseract."""
    text = _ocr_with_easyocr(pil_image_or_path)
    if text and len(text) > 20:
        return text

    if hasattr(pil_image_or_path, "size"):
        fallback = _ocr_with_tesseract(pil_image_or_path)
    else:
        try:
            from PIL import Image

            with Image.open(pil_image_or_path) as img:
                prepared = img.convert("RGB") if img.mode != "RGB" else img.copy()
            fallback = _ocr_with_tesseract(prepared)
        except Exception as e:
            logger.debug("Fallback OCR image open failed: %s", e)
            fallback = ""

    return fallback or text or ""


def extract_text_from_file(file_path: Path, ext: str) -> str:
    """
    Извлекает текст из файла.
    Поддерживает PDF, изображения, txt, docx.
    Для PDF и изображений используется OCR при необходимости.
    """
    if not file_path:
        return ""
    if not isinstance(file_path, Path):
        file_path = Path(file_path)
    if not file_path.exists() or not file_path.is_file():
        return ""

    ext = (ext or file_path.suffix or "").lower().lstrip(".")

    try:
        if ext == "pdf":
            return _extract_pdf(file_path)
        if ext in ("jpg", "jpeg", "png", "webp", "bmp", "tiff", "tif", "heic"):
            return _extract_image(file_path)
        if ext == "txt":
            return _extract_txt(file_path)
        if ext == "docx":
            return _extract_docx(file_path)
    except Exception as e:
        logger.warning("document_extraction failed path=%s ext=%s: %s", file_path, ext, e)

    return ""


def _extract_pdf(file_path: Path) -> str:
    """Текст из PDF: сначала встроенный текстовый слой; при малом тексте — OCR по страницам."""
    try:
        import fitz
    except ImportError:
        logger.warning("PyMuPDF not installed; PDF extraction disabled")
        return ""

    try:
        text_parts: list[str] = []
        with fitz.open(file_path) as doc:
            for page in doc:
                try:
                    text_parts.append(page.get_text() or "")
                except Exception as page_error:
                    logger.debug("PDF text layer page read error: %s", page_error)

        text_layer = _normalize_text("\n".join(part for part in text_parts if part))

        if len(text_layer) >= MIN_PDF_TEXT_FOR_OCR:
            return text_layer

        try:
            from PIL import Image
        except ImportError:
            return text_layer

        ocr_parts: list[str] = []
        with fitz.open(file_path) as doc:
            for page in doc:
                try:
                    try:
                        pix = page.get_pixmap(dpi=150, alpha=False)
                    except TypeError:
                        pix = page.get_pixmap(matrix=fitz.Matrix(2, 2), alpha=False)

                    try:
                        pil_img = pix.pil_image()
                    except (AttributeError, TypeError):
                        pil_img = Image.frombytes("RGB", (pix.width, pix.height), pix.samples)

                    page_text = _run_ocr(pil_img)
                    if page_text:
                        ocr_parts.append(page_text)
                except Exception as page_error:
                    logger.debug("PDF OCR page error: %s", page_error)

        ocr_text = _normalize_text("\n\n".join(ocr_parts))

        if len(ocr_text) > len(text_layer):
            return ocr_text
        return text_layer or ocr_text
    except Exception as e:
        logger.warning("PDF extraction error: %s", e)
        return ""


def _extract_txt(file_path: Path) -> str:
    """Чтение plain text с безопасной деградацией по кодировкам."""
    for encoding in ("utf-8", "utf-8-sig", "cp1251", "latin-1"):
        try:
            return _normalize_text(file_path.read_text(encoding=encoding, errors="replace"))
        except Exception:
            continue

    try:
        return _normalize_text(file_path.read_text(errors="replace"))
    except Exception as e:
        logger.warning("TXT read error: %s", e)
        return ""


def _extract_docx(file_path: Path) -> str:
    """Извлечение текста из .docx без ломки архитектуры."""
    try:
        import docx
    except ImportError:
        logger.debug("python-docx not installed; DOCX extraction skipped")
        return ""

    try:
        document = docx.Document(str(file_path))
        parts: list[str] = []

        for paragraph in document.paragraphs:
            if paragraph.text and paragraph.text.strip():
                parts.append(paragraph.text.strip())

        for table in document.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text and cell.text.strip()]
                if cells:
                    parts.append(" | ".join(cells))

        return _normalize_text("\n".join(parts))
    except Exception as e:
        logger.warning("DOCX extraction error: %s", e)
        return ""


def _extract_image(file_path: Path) -> str:
    """Изображение: EasyOCR, резерв — Tesseract."""
    try:
        from PIL import Image
    except ImportError:
        return ""

    try:
        with Image.open(file_path) as source_img:
            img = source_img.copy()

        if img.mode not in ("L", "RGB", "RGBA"):
            img = img.convert("RGB")
        elif img.mode == "RGBA":
            background = Image.new("RGB", img.size, (255, 255, 255))
            background.paste(img, mask=img.split()[3])
            img = background

        w, h = img.size
        max_side = 2000
        if max(w, h) > max_side:
            ratio = max_side / max(w, h)
            new_size = (max(1, int(w * ratio)), max(1, int(h * ratio)))
            resampling = getattr(Image, "Resampling", Image).LANCZOS
            img = img.resize(new_size, resampling)

        text = _run_ocr(img)
        if not text and file_path.exists():
            text = _run_ocr(file_path)

        return _normalize_text(text)
    except Exception as e:
        logger.warning("Image OCR error: %s", e)
        return ""